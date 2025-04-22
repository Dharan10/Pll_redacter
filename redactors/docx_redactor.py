import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from typing import Dict, List, Any, Optional, Union, Tuple
import logging
from pathlib import Path
import os
import re

logger = logging.getLogger(__name__)

class DocxRedactor:
    """Redacts PII from Microsoft Word (DOCX) documents.
    
    This class handles the redaction of PII from DOCX files using python-docx,
    with options for text replacement or highlighting.
    """
    
    def __init__(self, redaction_marker: str = "[REDACTED]", 
                 highlight_mode: bool = False,
                 highlight_color: WD_COLOR_INDEX = WD_COLOR_INDEX.BLACK):
        """Initialize the DOCX redactor.
        
        Args:
            redaction_marker: The text to replace PII with
            highlight_mode: If True, highlight PII instead of replacing it
            highlight_color: Color to use for highlighting
        """
        self.redaction_marker = redaction_marker
        self.highlight_mode = highlight_mode
        self.highlight_color = highlight_color
    
    def redact(self, docx_path: Union[str, Path], pii_list: List[Dict[str, Any]],
               output_path: Union[str, Path]) -> bool:
        """Redact PII from a DOCX document.
        
        Args:
            docx_path: Path to the input DOCX file
            pii_list: List of PII items with 'value' field
            output_path: Path to save the redacted DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Open the DOCX document
            doc = docx.Document(docx_path)
            
            # Extract PII values
            pii_values = [pii['value'] for pii in pii_list]
            
            # Process paragraphs in the document
            self._process_paragraphs(doc.paragraphs, pii_values)
            
            # Process tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._process_paragraphs(cell.paragraphs, pii_values)
            
            # Process headers and footers
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        self._process_paragraphs(header.paragraphs, pii_values)
                
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        self._process_paragraphs(footer.paragraphs, pii_values)
            
            # Save the redacted document
            doc.save(output_path)
            
            logger.info(f"Redacted DOCX saved to {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error redacting DOCX {docx_path}: {e}")
            return False
    
    def _process_paragraphs(self, paragraphs, pii_values: List[str]) -> None:
        """Process paragraphs to redact PII.
        
        Args:
            paragraphs: List of paragraph objects
            pii_values: List of PII values to redact
        """
        for paragraph in paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            # Check if paragraph contains any PII
            original_text = paragraph.text
            contains_pii = False
            
            for pii_value in pii_values:
                if pii_value in original_text:
                    contains_pii = True
                    break
            
            if not contains_pii:
                continue
            
            # Clear the paragraph
            paragraph.clear()
            
            # Process the text with PII redaction
            if self.highlight_mode:
                # Highlight mode: highlight the PII
                self._highlight_pii_in_text(paragraph, original_text, pii_values)
            else:
                # Replacement mode: replace PII with redaction marker
                redacted_text = self._replace_pii_in_text(original_text, pii_values)
                paragraph.add_run(redacted_text)
    
    def _replace_pii_in_text(self, text: str, pii_values: List[str]) -> str:
        """Replace PII values in text with redaction marker.
        
        Args:
            text: Original text
            pii_values: List of PII values to redact
            
        Returns:
            Redacted text
        """
        redacted_text = text
        
        # Sort PII values by length (descending) to handle overlapping values correctly
        sorted_pii = sorted(pii_values, key=len, reverse=True)
        
        for pii_value in sorted_pii:
            # Escape special regex characters in the PII value
            escaped_pii = re.escape(pii_value)
            # Replace the PII with the redaction marker
            redacted_text = re.sub(escaped_pii, self.redaction_marker, redacted_text)
        
        return redacted_text
    
    def _highlight_pii_in_text(self, paragraph, text: str, pii_values: List[str]) -> None:
        """Add text to paragraph with PII highlighted.
        
        Args:
            paragraph: Paragraph object to add text to
            text: Original text
            pii_values: List of PII values to highlight
        """
        # Find all PII positions in the text
        positions = []
        
        for pii_value in pii_values:
            for match in re.finditer(re.escape(pii_value), text):
                positions.append((match.start(), match.end(), pii_value))
        
        # Sort positions by start index
        positions.sort(key=lambda x: x[0])
        
        # Add text with highlighting
        last_end = 0
        
        for start, end, pii_value in positions:
            # Add text before PII
            if start > last_end:
                paragraph.add_run(text[last_end:start])
            
            # Add highlighted PII
            run = paragraph.add_run(pii_value)
            run.font.highlight_color = self.highlight_color
            
            last_end = end
        
        # Add remaining text after last PII
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def extract_text_with_positions(self, docx_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """Extract text with structural information from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of dictionaries containing text and structural information
        """
        try:
            # Open the DOCX document
            doc = docx.Document(docx_path)
            
            results = []
            
            # Extract text from paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    results.append({
                        'text': paragraph.text,
                        'type': 'paragraph',
                        'index': i,
                        'style': paragraph.style.name
                    })
            
            # Extract text from tables
            for t, table in enumerate(doc.tables):
                for r, row in enumerate(table.rows):
                    for c, cell in enumerate(row.cells):
                        cell_text = '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
                        if cell_text:
                            results.append({
                                'text': cell_text,
                                'type': 'table_cell',
                                'table_index': t,
                                'row_index': r,
                                'column_index': c
                            })
            
            return results
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
            return []
    
    def set_redaction_marker(self, marker: str) -> None:
        """Set the redaction marker text.
        
        Args:
            marker: The text to replace PII with
        """
        self.redaction_marker = marker
        logger.info(f"Redaction marker set to: {marker}")
    
    def set_highlight_mode(self, highlight_mode: bool) -> None:
        """Set the redaction mode.
        
        Args:
            highlight_mode: If True, highlight PII instead of replacing it
        """
        self.highlight_mode = highlight_mode
        mode_str = "highlight" if highlight_mode else "text replacement"
        logger.info(f"Redaction mode set to: {mode_str}")
    
    def set_highlight_color(self, color: WD_COLOR_INDEX) -> None:
        """Set the color for highlighting.
        
        Args:
            color: Color to use for highlighting
        """
        self.highlight_color = color
        logger.info(f"Highlight color set to: {color}")